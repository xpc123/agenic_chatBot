# -*- coding: utf-8 -*-
"""
文档处理器 - Document Processor
支持多种文档格式的解析和分块
"""
from typing import List, Dict, Any, Optional
from pathlib import Path
from loguru import logger
import hashlib
from datetime import datetime

from ..models.document import Document, DocumentChunk
from ..config import settings


class DocumentProcessor:
    """
    文档处理器
    
    职责:
    1. 解析多种文档格式 (PDF, DOCX, TXT, MD, HTML)
    2. 智能文档分块 (Chunking)
    3. 提取元数据
    """
    
    def __init__(self):
        self.chunk_size = settings.CHUNK_SIZE
        self.chunk_overlap = settings.CHUNK_OVERLAP
        
        logger.info("DocumentProcessor initialized")
    
    async def process_document(
        self,
        file_path: str,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> Document:
        """
        处理文档
        
        Args:
            file_path: 文档路径
            metadata: 额外元数据
        
        Returns:
            文档对象
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # 生成文档ID
        doc_id = self._generate_doc_id(file_path)
        
        # 提取文本
        text = await self._extract_text(file_path)
        
        # 分块
        chunks = await self._chunk_text(text, doc_id)
        
        # 创建文档对象
        document = Document(
            id=doc_id,
            filename=path.name,
            file_type=path.suffix,
            file_size=path.stat().st_size,
            processed=True,
            chunk_count=len(chunks),
            metadata=metadata or {},
        )
        
        logger.info(f"Processed document: {path.name}, {len(chunks)} chunks")
        
        return document, chunks
    
    async def _extract_text(self, file_path: str) -> str:
        """提取文档文本"""
        path = Path(file_path)
        extension = path.suffix.lower()
        
        if extension == '.pdf':
            return await self._extract_pdf(file_path)
        elif extension == '.docx':
            return await self._extract_docx(file_path)
        elif extension in ['.txt', '.md']:
            return await self._extract_text_file(file_path)
        elif extension == '.html':
            return await self._extract_html(file_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")
    
    async def _extract_pdf(self, file_path: str) -> str:
        """提取PDF文本"""
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(file_path)
            text = ""
            
            for page in reader.pages:
                text += page.extract_text() + "\n\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise
    
    async def _extract_docx(self, file_path: str) -> str:
        """提取DOCX文本"""
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(file_path)
            text = "\n\n".join([para.text for para in doc.paragraphs])
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise
    
    async def _extract_text_file(self, file_path: str) -> str:
        """提取纯文本"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Text extraction failed: {e}")
            raise
    
    async def _extract_html(self, file_path: str) -> str:
        """提取HTML文本"""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
            
            # 移除script和style标签
            for script in soup(["script", "style"]):
                script.decompose()
            
            return soup.get_text(separator='\n\n').strip()
            
        except Exception as e:
            logger.error(f"HTML extraction failed: {e}")
            raise
    
    async def _chunk_text(
        self,
        text: str,
        doc_id: str,
    ) -> List[DocumentChunk]:
        """
        智能文本分块
        
        使用滑动窗口和重叠来保持上下文连贯性
        """
        chunks = []
        
        # 简单的字符级分块 (可以优化为语义分块)
        text_length = len(text)
        chunk_index = 0
        
        for start in range(0, text_length, self.chunk_size - self.chunk_overlap):
            end = min(start + self.chunk_size, text_length)
            chunk_text = text[start:end].strip()
            
            if not chunk_text:
                continue
            
            chunk = DocumentChunk(
                id=f"{doc_id}_chunk_{chunk_index}",
                document_id=doc_id,
                content=chunk_text,
                chunk_index=chunk_index,
                metadata={
                    "start_char": start,
                    "end_char": end,
                    "length": len(chunk_text),
                }
            )
            
            chunks.append(chunk)
            chunk_index += 1
        
        return chunks
    
    def _generate_doc_id(self, file_path: str) -> str:
        """生成文档ID"""
        path = Path(file_path)
        content = f"{path.name}_{path.stat().st_mtime}"
        return hashlib.md5(content.encode()).hexdigest()[:16]
    
    async def chunk_with_semantic_splitting(
        self,
        text: str,
        doc_id: str,
        use_embeddings: bool = False,
    ) -> List[DocumentChunk]:
        """
        语义分块 (高级功能)
        
        基于段落、句子边界和语义相似度进行智能分块
        
        策略：
        1. 首先按段落分割
        2. 对于过长的段落，按句子边界分割
        3. 可选：使用 embedding 相似度合并相关的小段落
        
        Args:
            text: 原始文本
            doc_id: 文档ID
            use_embeddings: 是否使用 embedding 进行语义合并
        
        Returns:
            文档块列表
        """
        # 第一步：按段落分割
        paragraphs = self._split_by_paragraphs(text)
        
        # 第二步：处理每个段落
        semantic_chunks = []
        for para in paragraphs:
            if len(para) <= self.chunk_size:
                # 段落大小合适，直接作为一个 chunk
                semantic_chunks.append(para)
            else:
                # 段落太长，按句子边界分割
                sentences = self._split_by_sentences(para)
                semantic_chunks.extend(self._merge_sentences_to_chunks(sentences))
        
        # 第三步：可选 - 使用 embedding 合并相似的小段落
        if use_embeddings and len(semantic_chunks) > 1:
            semantic_chunks = await self._merge_by_semantic_similarity(semantic_chunks)
        
        # 第四步：创建 DocumentChunk 对象
        chunks = []
        for idx, chunk_text in enumerate(semantic_chunks):
            if not chunk_text.strip():
                continue
            
            chunk = DocumentChunk(
                id=f"{doc_id}_semantic_{idx}",
                document_id=doc_id,
                content=chunk_text.strip(),
                chunk_index=idx,
                metadata={
                    "chunking_method": "semantic",
                    "length": len(chunk_text),
                }
            )
            chunks.append(chunk)
        
        logger.info(f"Semantic chunking: {len(chunks)} chunks from {len(paragraphs)} paragraphs")
        return chunks
    
    def _split_by_paragraphs(self, text: str) -> List[str]:
        """按段落分割文本"""
        import re
        
        # 使用多种段落分隔符
        # 1. 双换行
        # 2. Markdown 标题
        # 3. 编号列表开头
        
        # 首先按双换行分割
        paragraphs = re.split(r'\n\s*\n', text)
        
        result = []
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # 检查是否是 Markdown 标题，如果是则单独作为一个 chunk
            if re.match(r'^#{1,6}\s+', para):
                result.append(para)
            # 检查是否包含多个 Markdown 标题
            elif re.search(r'\n#{1,6}\s+', para):
                # 按标题分割
                sub_parts = re.split(r'(?=\n#{1,6}\s+)', para)
                result.extend([p.strip() for p in sub_parts if p.strip()])
            else:
                result.append(para)
        
        return result
    
    def _split_by_sentences(self, text: str) -> List[str]:
        """按句子分割文本"""
        import re
        
        # 中英文句子分隔符
        # 中文：。！？；
        # 英文：. ! ? ; (后面跟空格或换行)
        
        sentence_pattern = r'([。！？；]|(?<=[.!?;])\s+)'
        
        parts = re.split(sentence_pattern, text)
        
        # 合并分隔符到前一个句子
        sentences = []
        current = ""
        
        for part in parts:
            if re.match(sentence_pattern, part):
                current += part
                if current.strip():
                    sentences.append(current.strip())
                current = ""
            else:
                current += part
        
        if current.strip():
            sentences.append(current.strip())
        
        return sentences
    
    def _merge_sentences_to_chunks(self, sentences: List[str]) -> List[str]:
        """将句子合并为合适大小的 chunk"""
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            # 如果加上这个句子不超过限制，就加上
            if len(current_chunk) + len(sentence) + 1 <= self.chunk_size:
                if current_chunk:
                    current_chunk += " " + sentence
                else:
                    current_chunk = sentence
            else:
                # 保存当前 chunk，开始新的
                if current_chunk:
                    chunks.append(current_chunk)
                
                # 如果单个句子就超过限制，强制分割
                if len(sentence) > self.chunk_size:
                    # 按字符分割，尽量在空格处断开
                    for i in range(0, len(sentence), self.chunk_size - self.chunk_overlap):
                        sub_chunk = sentence[i:i + self.chunk_size]
                        chunks.append(sub_chunk)
                    current_chunk = ""
                else:
                    current_chunk = sentence
        
        if current_chunk:
            chunks.append(current_chunk)
        
        return chunks
    
    async def _merge_by_semantic_similarity(
        self, 
        chunks: List[str], 
        similarity_threshold: float = 0.8
    ) -> List[str]:
        """
        使用 embedding 相似度合并相似的小段落
        
        只合并连续且相似的小段落，避免打乱顺序
        """
        if len(chunks) <= 1:
            return chunks
        
        try:
            from .embeddings import embedding_generator
            
            # 只处理小于 chunk_size/2 的段落
            small_threshold = self.chunk_size // 2
            
            # 生成 embeddings
            embeddings = await embedding_generator.embed_batch(chunks)
            
            # 合并相似的连续小段落
            merged_chunks = []
            current_merged = chunks[0]
            current_embedding = embeddings[0]
            
            for i in range(1, len(chunks)):
                chunk = chunks[i]
                embedding = embeddings[i]
                
                # 检查是否可以合并
                can_merge = (
                    len(current_merged) < small_threshold and
                    len(chunk) < small_threshold and
                    len(current_merged) + len(chunk) <= self.chunk_size and
                    self._cosine_similarity(current_embedding, embedding) >= similarity_threshold
                )
                
                if can_merge:
                    current_merged += "\n\n" + chunk
                    # 更新 embedding 为平均值
                    current_embedding = [
                        (a + b) / 2 for a, b in zip(current_embedding, embedding)
                    ]
                else:
                    merged_chunks.append(current_merged)
                    current_merged = chunk
                    current_embedding = embedding
            
            merged_chunks.append(current_merged)
            
            logger.debug(f"Semantic merging: {len(chunks)} -> {len(merged_chunks)} chunks")
            return merged_chunks
            
        except Exception as e:
            logger.warning(f"Semantic merging failed, using original chunks: {e}")
            return chunks
    
    def _cosine_similarity(self, vec1: List[float], vec2: List[float]) -> float:
        """计算余弦相似度"""
        import math
        
        dot_product = sum(a * b for a, b in zip(vec1, vec2))
        norm1 = math.sqrt(sum(a * a for a in vec1))
        norm2 = math.sqrt(sum(b * b for b in vec2))
        
        if norm1 == 0 or norm2 == 0:
            return 0.0
        
        return dot_product / (norm1 * norm2)
